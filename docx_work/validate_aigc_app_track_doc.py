from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "output" / "doc" / "win_应用赛道初赛作品策划.docx"

TERMS = [
    "2026年中国高校计算机大赛-AIGC创新赛",
    "作品名称：《智存》",
    "团队名称",
    "win",
    "李瑞麟",
    "黄隽尚",
    "罗子超",
    "石俊贤",
    "华中科技大学",
    "谭志虎",
    "作品概述",
    "作品宣传海报",
    "作品设计理念",
    "产品原型设计",
    "创新点说明",
    "前景评估",
    "大模型具体应用说明",
    "数据不出端",
]


def main() -> None:
    doc = Document(DOCX)
    text_parts = []
    text_parts.extend(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)
    all_text = "\n".join(text_parts)

    print(f"path={DOCX}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} sections={len(doc.sections)}")
    print(f"image_paragraphs={sum(1 for p in doc.paragraphs if p._p.xpath('.//w:drawing'))}")
    missing = []
    for term in TERMS:
        ok = term in all_text
        print(f"{term}: {'OK' if ok else 'MISSING'}")
        if not ok:
            missing.append(term)

    with ZipFile(DOCX) as zf:
        names = zf.namelist()
        media = [n for n in names if n.startswith("word/media/")]
        comments = [n for n in names if "comments" in n]
        doc_xml = zf.read("word/document.xml")
        core = etree.fromstring(zf.read("docProps/core.xml"))
        creator = ""
        last_modified_by = ""
        for el in core.iter():
            local = etree.QName(el).localname
            if local == "creator":
                creator = el.text or ""
            elif local == "lastModifiedBy":
                last_modified_by = el.text or ""
        print(f"media_count={len(media)}")
        print(f"comments_count={len(comments)}")
        print(f"creator={creator}")
        print(f"lastModifiedBy={last_modified_by}")
        print(f"track_revisions={'YES' if b'trackRevisions=\"1\"' in doc_xml or b'w:trackRevisions=\"1\"' in doc_xml else 'NO'}")

    if missing:
        raise SystemExit(f"Missing required terms: {missing}")
    if len(doc.tables) < 6:
        raise SystemExit("Expected at least 6 tables")
    if sum(1 for p in doc.paragraphs if p._p.xpath(".//w:drawing")) < 5:
        raise SystemExit("Expected original 5 image paragraphs to be preserved")


if __name__ == "__main__":
    main()
