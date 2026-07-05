from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "doc" / "win_字体统一黑色版.docx"
OUT = ROOT / "output" / "doc" / "win_应用赛道初赛作品策划.docx"


BLACK = RGBColor(0, 0, 0)
ACCENT = RGBColor(31, 78, 121)
LIGHT = "EAF2F8"
MID = "D9EAF7"
PALE = "F7FBFD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color=BLACK) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def format_paragraph(paragraph, size=10.5, before=0, after=6, line=1.15, align=None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size={1: 17, 2: 13, 3: 11.5}.get(level, 11), bold=True, color=BLACK)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    return p


def add_body(doc: Document, text: str, bold_prefix: str | None = None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_plain(doc: Document, text: str, size=10.5, bold=False, align=None, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=BLACK)
    return p


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def clear_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def remove_empty_paragraph_before(paragraph) -> None:
    prev = paragraph._p.getprevious()
    if prev is not None and prev.tag == qn("w:p") and not "".join(prev.xpath(".//w:t/text()")).strip() and not prev.xpath(".//w:drawing"):
        paragraph._p.getparent().remove(prev)


def append_body_elements(dst: Document, src: Document) -> None:
    body = dst.element.body
    sect_pr = body.sectPr
    for child in list(src.element.body):
        if child.tag == qn("w:sectPr"):
            continue
        if sect_pr is not None:
            body.insert(body.index(sect_pr), deepcopy(child))
        else:
            body.append(deepcopy(child))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        if name not in [s.name for s in doc.styles]:
            continue
        style = doc.styles[name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = BLACK
    doc.styles["Normal"].font.size = Pt(10.5)

    for p in doc.sections[0].header.paragraphs:
        p.text = ""
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_p.add_run("2026年中国高校计算机大赛-AIGC创新赛 · 应用赛道初赛作品策划")
    set_run_font(r, size=9, color=BLACK)

    for p in doc.sections[0].footer.paragraphs:
        p.text = ""
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("团队 win｜华中科技大学")
    set_run_font(r, size=9, color=BLACK)


def add_cover(doc: Document) -> None:
    add_plain(doc, "2026年中国高校计算机大赛-AIGC创新赛", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("应用赛道初赛作品策划")
    set_run_font(r, size=22, bold=True, color=BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("作品名称：《智存》")
    set_run_font(r, size=18, bold=True, color=BLACK)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = ["团队名称", "所在学校", "团队成员", "指导老师", "作品方向", "核心主题"]
    values = [
        "win",
        "华中科技大学",
        "李瑞麟、黄隽尚、罗子超、石俊贤",
        "谭志虎",
        "端侧AI与移动应用创新",
        "让每一次截图都成为可复用的知识资产",
    ]
    widths = [Cm(3.2), Cm(11.6)]
    for row, label, value in zip(table.rows, labels, values):
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 140, 180, 140, 180)
            set_cell_shading(cell, LIGHT if i == 0 else "FFFFFF")
        row.cells[0].text = label
        row.cells[1].text = value
        for p in row.cells[0].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_paragraph(p, size=11, after=0)
            for run in p.runs:
                run.bold = True
        for p in row.cells[1].paragraphs:
            format_paragraph(p, size=11, after=0)

    add_plain(doc, "", after=16)
    add_plain(doc, "提交说明", size=12, bold=True, after=4)
    add_body(
        doc,
        "本策划文档依据应用赛道初赛作品策划模板整理，围绕团队介绍、作品简介、作品策划三部分展开，重点说明作品的设计理念、产品原型、创新点与前景评估。",
    )
    add_page_break(doc)


def add_team_intro(doc: Document) -> None:
    add_heading(doc, "一、团队介绍", 1)
    add_heading(doc, "1. 团队基本信息", 2)
    table = doc.add_table(rows=7, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = [
        ("项目", "内容", "说明"),
        ("团队名称", "win", "参赛团队名称"),
        ("所在学校", "华中科技大学", "四位团队成员均来自华中科技大学"),
        ("李瑞麟", "产品经理 / 策划统筹", "负责需求洞察、产品定位、功能边界与文档统筹"),
        ("黄隽尚", "UI 与交互设计", "负责移动端信息架构、界面风格、核心流程与交互细节"),
        ("罗子超", "前端开发 / 原型实现", "负责静态H5原型、页面实现、演示路径与前端交互验证"),
        ("石俊贤", "服务端开发 / AI 工程", "负责数据结构、端侧AI接入方案、检索逻辑与工程可行性验证"),
    ]
    widths = [Cm(3.0), Cm(4.8), Cm(7.0)]
    for r_idx, (a, b, c) in enumerate(rows):
        row = table.rows[r_idx]
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_shading(cell, MID if r_idx == 0 else ("F7FBFD" if i == 0 else "FFFFFF"))
        for i, text in enumerate((a, b, c)):
            row.cells[i].text = text
            for p in row.cells[i].paragraphs:
                format_paragraph(p, size=10.2, after=0, align=WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
    set_repeat_table_header(table.rows[0])
    add_body(doc, "指导老师：谭志虎。")

    add_heading(doc, "2. 团队协作方式", 2)


def add_work_intro(doc: Document) -> None:
    add_heading(doc, "二、作品简介", 1)
    add_heading(doc, "1. 作品名称（必须）", 2)
    add_body(doc, "作品名称：《智存》。")
    add_heading(doc, "2. 作品概述（200字左右，必须）", 2)
    overview = (
        "智存是一款面向大学生、职场人士、游戏玩家和生活记录者的端侧AI知识管理应用，"
        "聚焦截图、速记、攻略、课程通知等碎片信息“保存容易、整理困难、检索低效”的高频痛点。"
        "作品以vivo蓝心端侧大模型为能力底座，通过全局悬浮球实现跨App无感捕获，"
        "利用OCR、多模态理解和自动打标将截图与文本沉淀为知识卡片，并以智能文件夹、瀑布流和语义搜索帮助用户秒级寻回。"
        "其突出创新在于全局伴随式入口、异构知识卡片重组、双路召回与状态保活、数据不出端的隐私保护。"
    )
    add_body(doc, overview)
    add_heading(doc, "3. 作品宣传海报（必须）", 2)
    add_poster(doc)
    add_page_break(doc)


def add_poster(doc: Document) -> None:
    table = doc.add_table(rows=6, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell, 160, 160, 160, 160)
            set_cell_shading(cell, PALE)

    title = table.cell(0, 0).merge(table.cell(0, 3))
    title.text = "智存"
    for p in title.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=26, bold=True, color=BLACK)

    subtitle = table.cell(1, 0).merge(table.cell(1, 3))
    subtitle.text = "让每一次截图都成为可复用的知识资产"
    for p in subtitle.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=14, bold=True, color=BLACK)

    features = [
        ("无感捕获", "悬浮球跨App长按截图，保存信息不打断当前心流。"),
        ("AI沉淀", "端侧大模型识别内容、生成摘要、标签与智能文件夹。"),
        ("秒级寻回", "关键词匹配与语义检索双路召回，返回列表保留状态。"),
        ("数据不出端", "截图进入私有沙盒，本地处理降低隐私泄露风险。"),
    ]
    for i, (name, desc) in enumerate(features):
        cell = table.cell(2, i)
        cell.text = ""
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(name)
        set_run_font(r, size=13, bold=True, color=BLACK)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(6)
        r = p2.add_run(desc)
        set_run_font(r, size=9.5, color=BLACK)
        set_cell_shading(cell, "EAF2F8")

    line = table.cell(3, 0).merge(table.cell(3, 3))
    line.text = "适用场景：考研复习｜游戏攻略｜课程通知｜旅行菜谱｜会议灵感"
    for p in line.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=11, bold=True, color=BLACK)

    flow = table.cell(4, 0).merge(table.cell(4, 3))
    flow.text = "捕获 → 理解 → 归档 → 搜索 → 复用"
    for p in flow.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=18, bold=True, color=BLACK)

    foot = table.cell(5, 0).merge(table.cell(5, 3))
    foot.text = "团队 win｜华中科技大学｜指导老师：谭志虎"
    for p in foot.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            set_run_font(run, size=10.5, color=BLACK)


def rename_existing_headings(doc: Document) -> None:
    mapping = {
        "团队介绍": "2. 团队定位与优势",
        "作品设计理念": "1. 作品设计理念",
        "作品原型设计": "2. 产品原型设计",
        "创新点": "3. 创新点说明",
        "前景评估": "4. 前景评估",
        "作品原型设计": "2. 产品原型设计",
    }
    for p in doc.paragraphs:
        text = p.text.strip()
        if text in mapping:
            p.text = mapping[text]
            p.style = "Heading 2" if text == "团队介绍" else "Heading 1"


def normalize_existing_content(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text and not p._p.xpath(".//w:drawing"):
            continue
        style = p.style.name if p.style is not None else ""
        if style == "Heading 1":
            size = 17
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(7)
            for run in p.runs:
                set_run_font(run, size=size, bold=True, color=BLACK)
        elif style == "Heading 2":
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(5)
            for run in p.runs:
                set_run_font(run, size=13, bold=True, color=BLACK)
        elif style == "Heading 3":
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                set_run_font(run, size=11.5, bold=True, color=BLACK)
        else:
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.18
            if text.startswith("图："):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = None
                for run in p.runs:
                    set_run_font(run, size=9.5, bold=False, color=BLACK)
            elif p._p.xpath(".//w:drawing"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.keep_with_next = True
                for run in p.runs:
                    set_run_font(run, size=10.5, color=BLACK)
            else:
                p.paragraph_format.first_line_indent = Cm(0.74)
                for run in p.runs:
                    set_run_font(run, size=10.5, color=BLACK)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        if table.rows:
            set_repeat_table_header(table.rows[0])
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_margins(cell, 100, 120, 100, 120)
                set_cell_shading(cell, MID if r_idx == 0 else "FFFFFF")
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.08
                    if r_idx == 0 or c_idx == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=9.3 if len(cell.text) > 55 else 9.8, bold=(r_idx == 0), color=BLACK)


def normalize_strategy_heading_levels(doc: Document) -> None:
    """Keep template's 策划 subsections visually under the main wrapper."""
    strategy_heads = {"1. 作品设计理念", "2. 产品原型设计", "3. 创新点说明", "4. 前景评估"}
    in_strategy = False
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "三、作品策划":
            in_strategy = True
            continue
        if not in_strategy:
            continue
        if text in strategy_heads:
            p.style = "Heading 2"
        elif p.style.name == "Heading 2":
            p.style = "Heading 3"


def build() -> None:
    src_doc = Document(SRC)
    rename_existing_headings(src_doc)
    normalize_existing_content(src_doc)

    out_doc = Document(SRC)
    clear_body(out_doc)
    configure_document(out_doc)

    add_cover(out_doc)
    add_team_intro(out_doc)
    append_body_elements(out_doc, src_doc)

    # Remove the empty paragraph that can appear between the newly inserted team
    # intro and the copied source body.
    for p in out_doc.paragraphs:
        if p.text.strip() == "2. 团队定位与优势":
            remove_empty_paragraph_before(p)
            break

    # Insert the作品简介 section after copied team intro and before作品策划.
    for p in list(out_doc.paragraphs):
        if p.text.strip() == "1. 作品设计理念":
            p.insert_paragraph_before("")
            marker = p.insert_paragraph_before("__WORK_INTRO_MARKER__")
            break
    else:
        marker = None

    if marker is not None:
        tmp = Document(SRC)
        clear_body(tmp)
        add_work_intro(tmp)
        body = out_doc.element.body
        marker_el = marker._p
        idx = body.index(marker_el)
        for child in list(tmp.element.body):
            if child.tag == qn("w:sectPr"):
                continue
            body.insert(idx, deepcopy(child))
            idx += 1
        body.remove(marker_el)

    # Add the作品策划 wrapper heading before design/prototype/innovation/prospect.
    for p in list(out_doc.paragraphs):
        if p.text.strip() == "1. 作品设计理念":
            wrapper = p.insert_paragraph_before("三、作品策划")
            wrapper.style = "Heading 1"
            for run in wrapper.runs:
                set_run_font(run, size=17, bold=True, color=BLACK)
            wrapper.paragraph_format.keep_with_next = True
            wrapper.paragraph_format.space_before = Pt(14)
            wrapper.paragraph_format.space_after = Pt(7)
            break

    # Ensure product prototype section explicitly mentions model application.
    for p in list(out_doc.paragraphs):
        if p.text.strip() == "原型总体说明":
            note = p.insert_paragraph_before(
                "核心功能解读：智存的核心功能卖点，是用全局悬浮入口把碎片信息快速捕获到私有沙盒，再用端侧大模型完成理解、归档与语义检索。"
            )
            format_paragraph(note, size=10.5)
            note.paragraph_format.first_line_indent = Cm(0.74)
            break
    for p in list(out_doc.paragraphs):
        if p.text.strip() == "核心流程设计":
            note = p.insert_paragraph_before(
                "大模型具体应用说明：作品将大模型用于OCR结果理解、多模态内容识别、摘要生成、自动标签、智能文件夹归类、向量语义检索与行动卡片结构化抽取，并优先采用vivo蓝心端侧能力完成本地处理。"
            )
            format_paragraph(note, size=10.5)
            note.paragraph_format.first_line_indent = Cm(0.74)
            break

    # Final global formatting pass for newly inserted content.
    normalize_strategy_heading_levels(out_doc)
    normalize_existing_content(out_doc)

    props = out_doc.core_properties
    props.author = "win"
    props.last_modified_by = "win"
    props.title = "智存 - 应用赛道初赛作品策划"
    props.subject = "2026年中国高校计算机大赛-AIGC创新赛"
    props.keywords = "AIGC创新赛; 应用赛道; 智存; win; 华中科技大学"
    props.comments = ""

    OUT.parent.mkdir(parents=True, exist_ok=True)
    out_doc.save(OUT)
    scrub_revision_and_personal_metadata(OUT)
    print(OUT)


def scrub_revision_and_personal_metadata(path: Path) -> None:
    tmp = path.with_suffix(".tmp.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                data = data.replace(b'w:trackRevisions="1"', b'w:trackRevisions="0"')
            zout.writestr(item, data)
    tmp.replace(path)


if __name__ == "__main__":
    build()
