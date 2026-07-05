from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor


SOURCE = Path("output/doc/win_内容表述优化版.docx")
OUT = Path("output/doc/win_字体统一黑色版.docx")
FONT_NAME = "Microsoft YaHei"
BLACK = "000000"


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_rpr_font_and_color(rpr):
    rfonts = ensure_child(rpr, "w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), FONT_NAME)

    color = ensure_child(rpr, "w:color")
    color.set(qn("w:val"), BLACK)
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        qattr = qn(attr)
        if qattr in color.attrib:
            del color.attrib[qattr]


def set_style_font(style):
    style.font.name = FONT_NAME
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    set_rpr_font_and_color(rpr)


def set_run_font(run):
    run.font.name = FONT_NAME
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    set_rpr_font_and_color(rpr)


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def text_signature(doc):
    items = [("p", p.text) for p in doc.paragraphs]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                items.append((f"t{ti}r{ri}c{ci}", cell.text))
    return items


def set_document_defaults(doc):
    styles_element = doc.styles.element
    doc_defaults = ensure_child(styles_element, "w:docDefaults")
    rpr_default = ensure_child(doc_defaults, "w:rPrDefault")
    rpr = ensure_child(rpr_default, "w:rPr")
    set_rpr_font_and_color(rpr)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document(str(SOURCE))
    before = text_signature(doc)

    set_document_defaults(doc)

    for style in doc.styles:
        if hasattr(style, "font"):
            set_style_font(style)

    for paragraph in all_paragraphs(doc):
        for run in paragraph.runs:
            set_run_font(run)

    doc.save(str(OUT))

    final_doc = Document(str(OUT))
    if text_signature(final_doc) != before:
        raise RuntimeError("Text content changed while unifying font/color.")

    print(f"saved: {OUT}")
    print(f"font: {FONT_NAME}")
    print("color: black")
    print("content check: unchanged")
    print(f"paragraphs: {len(final_doc.paragraphs)}, tables: {len(final_doc.tables)}, images: {len(final_doc.inline_shapes)}")


if __name__ == "__main__":
    main()
