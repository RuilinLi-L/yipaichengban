from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET
from io import BytesIO
import argparse
import base64
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
PRIMARY_DOCX = OUT_DIR / "智存复赛四人团队分工与项目完善说明.docx"
FALLBACK_DOCX = OUT_DIR / "zhicun_fusai_team_division.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 36, 52)
MUTED = RGBColor(91, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GREEN = "E8F4EF"
PALE_GOLD = "FFF4D8"
BORDER = "C9D3DF"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(cell, text, size=9.0, bold=False, color=INK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.18
    for idx, part in enumerate(str(text).split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(part)
        set_run_font(run, size=size, bold=bold, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_in):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    existing_grid = tbl.tblGrid
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = int(round(widths_in[idx] * 1440))
            cell.width = Inches(widths_in[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width_dxa))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE, body_fill=WHITE, font_size=8.6):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, header_fill)
        set_cell_text(cell, header, size=8.8, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            shade_cell(cells[idx], body_fill)
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 and len(str(value)) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[idx], value, size=font_size, color=INK, align=align)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    return table


def add_para(doc, text="", size=10.5, bold=False, color=INK, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_run_font(run, size=13, bold=True, color=BLUE)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
    else:
        set_run_font(run, size=12, bold=True, color=DARK_BLUE)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_border(cell, color="D7DEE8", size="6")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.22
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_check_items(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, size=10.0, color=INK)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def set_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    r = hp.add_run("智存复赛分工说明 | 2026 AIGC创新赛 应用赛道")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    r = fp.add_run("内部执行稿")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(doc):
    add_para(doc, "2026 中国高校计算机大赛 - AIGC创新赛 / 应用赛道复赛", size=10.5, bold=True, color=MUTED, after=8)
    add_para(doc, "智存复赛四人团队分工与项目完善说明", size=24, bold=True, color=INK, after=4)
    add_para(doc, "交付导向版本：围绕复赛提交物、评分维度和 vivo 终端落地倒推任务。", size=12.0, color=MUTED, after=14)
    add_table(
        doc,
        ["字段", "内容"],
        [
            ("作品名称", "智存"),
            ("分工口径", "交付导向：每个人绑定明确提交物、评分项和验收标准"),
            ("适用团队", "四人团队，成员以 A-D 占位，可直接替换真实姓名"),
            ("当前节点", "2026-07-01 起执行，优先保证 2026-07-06 前完成复赛作品提交"),
            ("依据材料", "应用赛道说明、复赛提交清单、策划 PPT 模板、现有 Web/uni-app 原型与 AIAdapter 代码"),
        ],
        [1.35, 5.15],
        header_fill=LIGHT_BLUE,
        font_size=9.2,
    )
    add_callout(
        doc,
        "总原则",
        "复赛不是平均分工，而是交付物负责制。每位成员都要有主责产出、可验收标准和备用协作人；所有材料统一围绕“无感捕获、有序沉淀、秒级寻回”和“端侧隐私知识资产”表达。",
        fill=PALE_GREEN,
    )


def add_scoring_summary(doc):
    add_heading(doc, "一、赛事要求与评分导向摘要", 1)
    add_para(
        doc,
        "应用赛道要求作品能够落地 vivo 终端设备，作品形态可以是 APP、智能体、快应用或插件。复赛阶段开放云测平台与蓝心端侧大模型能力，因此本项目要把“能运行、能演示、能说明大模型调用、能解释端侧化路线”作为最低交付线。",
    )
    add_table(
        doc,
        ["复赛提交物", "必须做到", "主责成员", "验收标准"],
        [
            ("策划 PPT", "覆盖团队介绍、作品简介、设计理念、原型效果、创新点、前景评估、大模型说明", "成员 A", "按模板栏目完整填写，PPT 中出现核心界面、流程和代码/接口说明"),
            ("宣传海报", "突出作品名称、核心卖点、目标场景和端侧隐私优势", "成员 B", "一张可提交海报图，视觉与 PPT/演示视频统一"),
            ("演示视频", "约 3 分钟，竖版 9:16，展示产品功能与操作流程", "成员 B", "能看清 vivo 端或移动端操作，覆盖捕获、行动卡、知识库、搜索"),
            ("可运行产品包/链接", "提供 uni-app/Android 可运行版本或稳定演示链接", "成员 C", "在 vivo 真机或云测环境完成运行截图/录屏，保留打包说明"),
            ("大模型代码包", "提交核心功能调用大模型的代码与说明", "成员 D", "包含 AIAdapter、接口 schema、示例输入输出、隐私与端侧化说明"),
        ],
        [1.1, 2.15, 0.9, 2.35],
        font_size=8.2,
    )
    add_table(
        doc,
        ["评分维度", "复赛表达重点", "项目对应打法"],
        [
            ("作品创新性", "选题真实、理念贯穿、功能原创、界面完整、交互流畅", "用“截图不进相册、跨 App 无感捕获、异构知识卡片、状态保活搜索”支撑创新性"),
            ("应用价值", "技术可行、可转化、用户需求广泛且高频、市场趋势吻合", "聚焦学生学习、攻略收藏、报名通知、课程群消息等高频碎片化信息场景"),
            ("作品完成度", "资料齐全、原型或 demo 能运行、呈现质量高", "把 Web 原型作为展示兜底，把 uni-app/APK 作为复赛运行主线"),
            ("大模型应用能力", "说明大模型应用场景，可通过 API 调用实现产品功能", "展示 AIAdapter 将截图转为标题、摘要、标签、待办、材料和提醒的结构化输出"),
        ],
        [1.2, 2.3, 3.0],
        header_fill=LIGHT_BLUE,
        font_size=8.6,
    )


def add_principles(doc):
    add_heading(doc, "二、团队总体协作原则", 1)
    add_check_items(
        doc,
        [
            "统一作品名称：所有 PPT、海报、视频、代码说明和演示界面统一使用“智存”，避免“智存/一拍成办/内部代号”混用。",
            "统一核心叙事：作品不是普通截图管理工具，而是面向 vivo 端侧 AI 的个人碎片化知识管家。",
            "统一演示链路：开启捕获入口、导入或捕获截图、生成 AI 行动卡、保存到私有知识库、语义搜索找回、说明端侧化路线。",
            "统一提交节奏：先保复赛可提交，再做决赛增强；复赛优先主动导入和样例捕获，决赛再补真实 MediaProjection、Overlay Service 和模型微调。",
            "统一验收方式：每个任务必须产出文件、截图、录屏、代码或文字说明，不能只停留在口头完成。",
        ],
    )


def add_roles(doc):
    add_heading(doc, "三、四位成员详细职责、阶段任务与交付物", 1)
    add_table(
        doc,
        ["成员", "角色定位", "主责评分项", "核心产出"],
        [
            ("成员 A", "产品策划与参赛统筹", "应用价值、创新性、资料质量", "复赛 PPT 内容总控、痛点与前景评估、答辩逻辑、提交清单"),
            ("成员 B", "UI/交互与宣传表达", "界面、交互、资料呈现", "页面优化建议、流程图、宣传海报、3 分钟竖版视频"),
            ("成员 C", "移动端工程与 vivo 运行落地", "作品完成度、技术可行性", "uni-app 工程完善、vivo 运行、APK/链接、操作录屏"),
            ("成员 D", "大模型能力与端侧化方案", "大模型应用能力、隐私可信", "AIAdapter、OCR/多模态说明、代码包、蓝心端侧接入路线"),
        ],
        [0.85, 1.35, 1.45, 2.85],
        font_size=8.6,
    )

    role_blocks = [
        (
            "成员 A：产品策划与参赛统筹",
            "把项目从“有原型”整理成“评委能快速理解且愿意给分”的参赛作品。",
            [
                ("复赛前必须完成", "整理 PPT 主线：团队介绍、作品简介、设计理念、核心创新点、前景评估；统一所有文案里的作品名称和核心卖点；建立最终提交清单并每天检查缺口。"),
                ("每天要做什么", "上午同步工程/AI/设计进度，更新一页提交看板；下午把新增功能截图和说明补进 PPT；晚上按评分表逐项查漏补缺。"),
                ("交付物", "复赛策划 PPT 终稿、答辩讲稿提纲、提交清单核对表、项目一页式简介。"),
                ("验收标准", "PPT 不缺模板必填项；每个评分维度至少有 2 个可展示证据；材料能在 3 分钟内讲清楚“痛点-方案-模型-落地”。"),
            ],
        ),
        (
            "成员 B：UI/交互与宣传表达",
            "把智存的体验做得可看、可讲、可传播，确保评委能在短时间内看到产品闭环。",
            [
                ("复赛前必须完成", "梳理 1 条主流程和 2 个辅助场景；优化首页、捕获页、行动卡、知识库、搜索页的展示截图；制作宣传海报；完成 9:16 竖版演示视频。"),
                ("每天要做什么", "上午确认当日可录制功能；下午补界面截图和流程图；晚上剪辑视频并同步给成员 A 更新 PPT。"),
                ("交付物", "海报图、演示视频 mp4、核心流程图、关键界面截图包、视频旁白稿。"),
                ("验收标准", "视频约 3 分钟、竖版 9:16、无明显卡顿或看不清文字；海报第一眼能读到“智存”和三段式卖点。"),
            ],
        ),
        (
            "成员 C：移动端工程与 vivo 运行落地",
            "把当前原型推进到复赛可运行状态，让“完成度”不被扣在工程落地上。",
            [
                ("复赛前必须完成", "以 uni-app 工程为运行主线，确认 HBuilderX 导入、依赖、页面路由、静态资源和本地存储可用；准备 vivo 真机或云测运行截图；保留 Web demo 作为兜底链接。"),
                ("每天要做什么", "上午修复运行问题；下午打包或真机测试；晚上输出运行说明、问题清单和可录制版本给成员 B。"),
                ("交付物", "可运行产品包或演示链接、HBuilderX 运行说明、vivo 真机/云测截图、关键功能录屏、工程 README。"),
                ("验收标准", "首页、捕获、知识金库、语义搜索、详情页、技术路线页都能走通；返回搜索列表后保留查询状态；本地数据可保存和恢复。"),
            ],
        ),
        (
            "成员 D：大模型能力与端侧化方案",
            "证明作品确实用了大模型，并且为 vivo 蓝心端侧能力预留了清晰接口。",
            [
                ("复赛前必须完成", "整理 AIAdapter 的输入输出协议；准备一次真实或可复现的 API 调用示例；说明 OCR、多模态理解、自动打标、行动卡生成、语义检索如何接入蓝心端侧模型。"),
                ("每天要做什么", "上午确认接口是否可调用；下午更新核心代码包和示例 JSON；晚上给成员 A/B 提供代码截图、模型流程图和视频讲解点。"),
                ("交付物", "大模型调用核心代码包、接口 schema、示例输入输出、端侧化技术说明、隐私边界说明。"),
                ("验收标准", "评委能看到“截图输入-模型理解-结构化行动卡输出”的闭环；代码包能定位到核心 API 调用和模型输出字段；不把端侧化写成空泛口号。"),
            ],
        ),
    ]
    for title, mission, rows in role_blocks:
        add_heading(doc, title, 2)
        add_callout(doc, "任务定位", mission, fill=LIGHT_GRAY)
        add_table(doc, ["项目", "说明"], rows, [1.25, 5.25], header_fill=LIGHT_BLUE, font_size=8.9)


def add_improvement_plan(doc):
    add_heading(doc, "四、项目当前状态与完善修改建议", 1)
    add_table(
        doc,
        ["当前已有基础", "复赛要补强的地方", "负责人"],
        [
            ("已有 Web 演示和 uni-app 工程骨架", "统一作品名为“智存”；复赛以 uni-app/vivo 运行为主，Web 演示作为兜底展示", "成员 C"),
            ("已有 AIAdapter 与 OpenAI 代理代码", "准备可提交代码包，补示例输入输出、接口 schema、模型失败回退和端侧替换说明", "成员 D"),
            ("已有私有知识库、样例捕获、语义搜索模拟", "把演示主线固定成“捕获-行动卡-知识库-搜索-技术路线”，减少分散功能介绍", "成员 A/B"),
            ("已有端侧化叙事", "把“应用私有目录、SQLite/KV、OCR/蓝心端侧、Overlay、MediaProjection”拆成复赛/决赛两阶段路线", "成员 C/D"),
            ("已有初赛策划文档与模板 PPT", "按复赛模板重写材料，补可运行版本、大模型代码、海报、竖版视频证据", "成员 A/B"),
        ],
        [2.05, 3.55, 0.9],
        header_fill=LIGHT_BLUE,
        font_size=8.5,
    )

    add_heading(doc, "复赛版本建议锁定的功能边界", 2)
    add_table(
        doc,
        ["模块", "复赛必须稳定", "可放到决赛继续增强"],
        [
            ("捕获入口", "主动上传截图、样例捕获、悬浮球模拟入口", "真实 Android Overlay Service、长按悬浮球跨 App 截屏"),
            ("AI 理解", "截图或文字生成标题、摘要、标签、待办、材料、提醒", "蓝心端侧模型离线多模态理解、模型微调"),
            ("知识库", "私有知识库、本地保存、智能文件夹、异构卡片展示", "SQLite/KV、向量索引持久化、批量导入"),
            ("搜索", "关键词匹配与语义模拟，返回后状态保活", "真实 embedding 检索、多模态相似度召回"),
            ("隐私说明", "不写入系统相册、不自动读取聊天/短信/通知、用户主动导入", "系统级权限边界、端侧安全沙箱与权限弹窗细化"),
        ],
        [1.15, 2.65, 2.7],
        header_fill=PALE_GREEN,
        font_size=8.4,
    )


def add_submission_checklist(doc):
    add_heading(doc, "五、复赛提交物清单对照表", 1)
    add_table(
        doc,
        ["序号", "提交物", "文件/内容建议命名", "责任人", "完成标准"],
        [
            ("1", "应用赛道复赛作品策划 PPT", "智存_复赛作品策划.pptx", "成员 A", "模板栏目全部完整；插入成员分工、核心界面、模型调用、端侧路线和提交说明"),
            ("2", "作品宣传海报图", "智存_宣传海报.png", "成员 B", "竖版或适合平台展示；包含作品名、三段式卖点、目标用户、vivo 端侧 AI 标签"),
            ("3", "作品演示视频", "智存_复赛演示视频_9比16.mp4", "成员 B", "约 3 分钟；按主流程录制；片尾出现作品名称和团队名称占位"),
            ("4", "可运行产品包或链接", "智存_uniapp工程包.zip / APK / 演示链接", "成员 C", "能在 vivo 真机、云测或本地移动端稳定打开；附运行说明"),
            ("5", "核心功能调用大模型代码包", "智存_AIAdapter_核心代码包.zip", "成员 D", "包含接口调用代码、schema、示例输入输出、端侧替换说明和隐私说明"),
        ],
        [0.5, 1.35, 1.75, 0.75, 2.15],
        header_fill=LIGHT_BLUE,
        font_size=8.0,
    )
    add_callout(
        doc,
        "提交前最后检查",
        "压缩包内不要只放源码而没有说明；PPT、视频、海报、代码包中作品名称必须一致；视频里展示的功能必须能在可运行版本中找到对应入口。",
        fill=PALE_GOLD,
    )


def add_schedule_and_risks(doc):
    add_heading(doc, "六、时间推进建议与风险预案", 1)
    add_para(doc, "以下排期以 2026-07-01 为起点，目标是在 2026-07-06 复赛作品提交窗口结束前完成可提交版本。")
    add_table(
        doc,
        ["日期", "成员 A", "成员 B", "成员 C", "成员 D", "当天验收"],
        [
            ("07-01", "定稿叙事框架和 PPT 目录", "确定视觉风格和视频脚本", "跑通 uni-app/确认运行问题", "整理 AIAdapter 与 schema", "四人确认主流程和功能边界"),
            ("07-02", "填完作品简介/理念/创新点", "补核心界面截图和流程图", "修复首页、知识库、搜索链路", "跑通或模拟 API 示例输出", "可演示版本第一次冻结"),
            ("07-03", "补前景评估和评分映射", "完成海报初稿", "准备真机/云测截图", "完成代码包说明初稿", "PPT 初稿 + 代码包初稿"),
            ("07-04", "整合所有材料进 PPT", "录制 9:16 视频素材", "输出产品包或链接", "补端侧化路线图", "视频粗剪 + 运行包可打开"),
            ("07-05", "全量审稿和提交清单核对", "完成视频终剪和海报终稿", "修复最后阻塞问题", "压缩核心代码包并检查路径", "全量材料预提交检查"),
            ("07-06", "最终提交与备份", "提供最终海报/视频", "提供最终运行包/链接", "提供最终代码包", "提交完成后保存云盘备份"),
        ],
        [0.65, 1.15, 1.15, 1.15, 1.15, 1.25],
        header_fill=LIGHT_BLUE,
        font_size=7.35,
    )
    add_table(
        doc,
        ["风险", "触发信号", "预案", "负责人"],
        [
            ("vivo 真机运行不稳定", "APK 打包失败、页面白屏、权限异常", "保留 H5 演示链接；视频中说明当前为 uni-app 可迁移版本；同时提交运行说明和问题边界", "成员 C"),
            ("大模型 API 不可用", "网络、Key、额度或模型接口失败", "保留本地 mock 输出；提供一次录屏或 JSON 示例；文档中明确 AIAdapter 可替换为蓝心端侧模型", "成员 D"),
            ("视频超过 3 分钟或重点分散", "剪辑后信息过多、讲不完", "按主流程删减，只保留捕获、行动卡、知识库、搜索、端侧说明五段", "成员 B"),
            ("PPT 材料不齐", "缺海报、代码截图、运行截图或流程图", "成员 A 维护提交清单，每天晚上逐项点名缺口并指定补交时间", "成员 A"),
            ("作品名称和表达不统一", "界面或文件仍出现旧名", "提交前全文搜索旧名；说明早期工程代号仅为历史命名，正式作品统一为智存", "成员 A/C"),
        ],
        [1.1, 1.45, 3.25, 0.7],
        header_fill=PALE_GOLD,
        font_size=8.3,
    )


def add_final_acceptance(doc):
    add_heading(doc, "七、最终验收口径", 1)
    add_check_items(
        doc,
        [
            "提交清单五项材料全部存在，并且文件名能一眼识别作品和用途。",
            "四位成员各自有明确主责交付物，PPT 团队介绍页能直接引用本说明中的分工。",
            "演示流程可以连续完成：捕获入口、AI 行动卡、确认保存、知识库沉淀、语义搜索、技术路线说明。",
            "大模型说明不是概念描述，至少出现核心代码路径、接口 schema、示例输入输出或调用截图。",
            "端侧化说明分清复赛已实现、复赛模拟、决赛计划，避免夸大当前完成度。",
            "所有材料统一突出：无感捕获、有序沉淀、秒级寻回、隐私端侧化。",
        ],
    )


def audit_docx_package(path):
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    texts = "".join(t.text or "" for t in root.findall(".//w:t", ns))
    required = [
        "智存",
        "成员 A",
        "成员 B",
        "成员 C",
        "成员 D",
        "策划 PPT",
        "宣传海报",
        "演示视频",
        "可运行产品包",
        "大模型代码包",
        "作品创新性",
        "应用价值",
        "作品完成度",
        "大模型应用能力",
    ]
    missing = [item for item in required if item not in texts]
    if missing:
        raise RuntimeError(f"Missing required text: {missing}")
    return {"chars": len(texts), "tables": len(root.findall('.//w:tbl', ns))}


def audit_docx_bytes(blob):
    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    with ZipFile(BytesIO(blob)) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    texts = "".join(t.text or "" for t in root.findall(".//w:t", ns))
    required = [
        "智存",
        "成员 A",
        "成员 B",
        "成员 C",
        "成员 D",
        "策划 PPT",
        "宣传海报",
        "演示视频",
        "可运行产品包",
        "大模型代码包",
        "作品创新性",
        "应用价值",
        "作品完成度",
        "大模型应用能力",
    ]
    missing = [item for item in required if item not in texts]
    if missing:
        raise RuntimeError(f"Missing required text: {missing}")
    return {"chars": len(texts), "tables": len(root.findall('.//w:tbl', ns))}


def make_document():
    doc = Document()
    style_document(doc)
    set_header_footer(doc)
    add_cover(doc)
    add_scoring_summary(doc)
    add_principles(doc)
    add_roles(doc)
    add_improvement_plan(doc)
    add_submission_checklist(doc)
    add_schedule_and_risks(doc)
    add_final_acceptance(doc)
    return doc


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = make_document()
    try:
        final_docx = PRIMARY_DOCX
        doc.save(final_docx)
    except PermissionError:
        final_docx = FALLBACK_DOCX
        doc.save(final_docx)
    audit = audit_docx_package(final_docx)
    print(final_docx)
    print(audit)


def emit_base64():
    doc = make_document()
    stream = BytesIO()
    doc.save(stream)
    blob = stream.getvalue()
    audit = audit_docx_bytes(blob)
    sys.stderr.write(str(audit) + "\n")
    sys.stdout.write(base64.b64encode(blob).decode("ascii"))


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--emit-base64", action="store_true")
    args = parser.parse_args()
    if args.emit_base64:
        emit_base64()
    else:
        build()
