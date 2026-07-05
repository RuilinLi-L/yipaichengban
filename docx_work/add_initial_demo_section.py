from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path("output/doc/win_排版完善版.docx")
OUT = Path("output/doc/win_补充初步功能演示版.docx")

IMG_CHAT = Path(
    r"C:\Users\RuilinLi\Documents\Tencent Files\1511964136\nt_qq\nt_data\Pic\2026-05\Thumb\dd2054b80fe2ea78c23bc5f71169817f_720.png"
)
IMG_CAPTURE = Path(
    r"C:\Users\RuilinLi\Documents\Tencent Files\1511964136\nt_qq\nt_data\Pic\2026-05\Ori\19e95bb3d3c3586507eb10822488d13c.png"
)
IMG_ACTION_TOP = Path(
    r"C:\Users\RuilinLi\Documents\Tencent Files\1511964136\nt_qq\nt_data\Pic\2026-05\Ori\5f42c9db66f92e8548317c48c8a2c250.png"
)
IMG_ACTION_BOTTOM = Path(
    r"C:\Users\RuilinLi\Documents\Tencent Files\1511964136\nt_qq\nt_data\Pic\2026-05\Ori\341c9e55856142515adfe4cda540138d.png"
)
IMG_KB = Path(
    r"C:\Users\RuilinLi\Documents\Tencent Files\1511964136\nt_qq\nt_data\Pic\2026-05\Ori\dd7057eb6f5de39d91af378114a233ce.png"
)

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
SUB_ACCENT = RGBColor(0x2F, 0x68, 0x92)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
GRID = "A7B9C9"
HEADER_FILL = "D9EAF7"
NOTE_FILL = "EDF7ED"


def set_rfonts(element, font_name=BODY_FONT):
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def format_runs(paragraph, size_pt=10.5, bold=None, color=TEXT):
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        set_rfonts(run._element, BODY_FONT)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        run.font.color.rgb = color


def format_body(paragraph, first_line=True, size_pt=10.5):
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.32
    pf.first_line_indent = Cm(0.74 if first_line else 0)
    pf.left_indent = Cm(0)
    format_runs(paragraph, size_pt=size_pt, bold=False, color=TEXT)


def format_heading(paragraph, level):
    paragraph.style = f"Heading {level}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.keep_with_next = True
    pf.keep_together = True
    pf.left_indent = Cm(0 if level < 3 else 0.35)
    pf.first_line_indent = Cm(0)
    if level == 2:
        pf.space_before = Pt(11)
        pf.space_after = Pt(5)
        pf.line_spacing = 1.15
        format_runs(paragraph, 13, bold=True, color=SUB_ACCENT)
    else:
        pf.space_before = Pt(7)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.15
        format_runs(paragraph, 11.5, bold=True, color=SUB_ACCENT)


def add_or_update_child(parent, child_name):
    child = parent.find(qn(child_name))
    if child is None:
        child = OxmlElement(child_name)
        parent.append(child)
    return child


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = add_or_update_child(tc_pr, "w:tcMar")
    for side, value in {
        "w:top": top,
        "w:start": start,
        "w:bottom": bottom,
        "w:end": end,
    }.items():
        node = mar.find(qn(side))
        if node is None:
            node = OxmlElement(side)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = add_or_update_child(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = add_or_update_child(tc_pr, "w:tcBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        edge = borders.find(qn(side))
        if edge is None:
            edge = OxmlElement(side)
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_table_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_fixed(table)
    set_repeat_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.18
                format_runs(p, 9.2 if row_idx else 9.7, bold=(row_idx == 0), color=ACCENT if row_idx == 0 else TEXT)


def add_note_table(doc, anchor, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    set_cell_border(cell, color="BBD7BD", size="6")
    set_cell_shading(cell, NOTE_FILL)
    p = cell.paragraphs[0]
    p.text = text
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.24
    format_runs(p, 10.2, bold=False, color=MUTED)
    anchor._p.addprevious(table._tbl)


def insert_caption(anchor, text):
    p = anchor.insert_paragraph_before(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.12
    format_runs(p, 9, bold=False, color=MUTED)
    return p


def insert_image(anchor, image_path, width_cm):
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return p


def add_flow_table(doc, anchor):
    rows = [
        ("环节", "用户动作", "原型反馈", "验证重点"),
        ("1. 导入", "将已做隐私处理的聊天记录截图上传，或在捕获页点击“演示报名通知”。", "截图进入“私有截图沙盒”，不写入系统相册。", "验证信息捕获与去相册污染。"),
        ("2. 抽取", "触发前端演示中的OCR与大模型结构化抽取。", "生成“确认专选课退选或补选安排”行动卡，并给出摘要、标签和建议。", "验证通知类截图可转为可行动信息。"),
        ("3. 确认", "进入行动卡页面，检查标题、摘要、待办、材料清单、知识标签等字段。", "对不确定的时间、地点保留“待确认”，由用户再核对。", "验证人机协同，避免误识别直接落库。"),
        ("4. 保存", "点击“确认保存”。", "卡片进入私有知识库，形成可检索的截图卡片。", "验证本地沉淀与标签检索闭环。"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    format_table(table, [2.1, 5.0, 5.1, 4.0])
    anchor._p.addprevious(table._tbl)


def base_signature(doc):
    items = [("p", p.text) for p in doc.paragraphs]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                items.append(("cell", cell.text))
    return items


def paragraph_texts(doc):
    return [p.text for p in doc.paragraphs]


def is_subsequence(needle, haystack):
    pos = 0
    for text in haystack:
        if pos < len(needle) and text == needle[pos]:
            pos += 1
    return pos == len(needle)


def main():
    missing = [p for p in [BASE, IMG_CHAT, IMG_CAPTURE, IMG_ACTION_TOP, IMG_ACTION_BOTTOM, IMG_KB] if not p.exists()]
    if missing:
        raise FileNotFoundError("\n".join(str(p) for p in missing))

    doc = Document(str(BASE))
    before = base_signature(doc)
    before_paragraphs = paragraph_texts(doc)
    anchors = [p for p in doc.paragraphs if p.text.strip() == "创新点"]
    if not anchors:
        raise RuntimeError("Could not find insertion point: 创新点")
    anchor = anchors[0]

    p = anchor.insert_paragraph_before("初步功能演示：聊天记录生成行动卡片")
    format_heading(p, 2)

    intro = anchor.insert_paragraph_before(
        "为验证“截图私有沙盒 → AI抽取 → 人工确认 → 本地沉淀”的闭环，原型新增了一个通知类场景演示：用户将已经做过隐私处理的课程群聊天记录作为输入，系统识别其中的退选、补选、课程调整等关键信息，并生成一张可确认、可保存、可检索的行动卡片。"
    )
    format_body(intro)

    p = anchor.insert_paragraph_before("演示流程")
    format_heading(p, 3)
    add_flow_table(doc, anchor)

    p = anchor.insert_paragraph_before("演示材料与界面状态")
    format_heading(p, 3)
    explain = anchor.insert_paragraph_before(
        "下列截图展示了从聊天记录输入到知识库卡片沉淀的完整链路：先导入已脱敏的聊天记录，再在捕获页触发演示，随后进入行动卡确认界面，最后保存为私有知识库中的一张截图卡片。"
    )
    format_body(explain)

    insert_caption(anchor, "图：已做隐私处理的聊天记录输入样例")
    insert_image(anchor, IMG_CHAT, 9.4)

    insert_caption(anchor, "图：捕获页提供上传截图与演示入口")
    insert_image(anchor, IMG_CAPTURE, 14.6)

    insert_caption(anchor, "图：行动卡上半部分展示标题、摘要与基础字段")
    insert_image(anchor, IMG_ACTION_TOP, 14.6)

    insert_caption(anchor, "图：行动卡下半部分支持标签、建议与确认保存")
    insert_image(anchor, IMG_ACTION_BOTTOM, 14.6)

    insert_caption(anchor, "图：保存后在私有知识库中生成可检索卡片")
    insert_image(anchor, IMG_KB, 14.6)

    add_note_table(
        doc,
        anchor,
        "说明：当前前端仅用于展示核心交互与数据流，是初步功能演示，不代表最终产品的视觉品质、动效完整度、模型能力和工程完成度。正式版本仍需继续完善UI精度、权限边界、端侧模型接入、异常提示、隐私合规与真实设备适配。",
    )

    doc.save(str(OUT))

    final_doc = Document(str(OUT))
    final_sig = base_signature(final_doc)
    if not is_subsequence(before_paragraphs, paragraph_texts(final_doc)):
        raise RuntimeError("Original paragraph order was not preserved.")
    for item in before:
        if item not in final_sig:
            raise RuntimeError(f"Existing content missing after insertion: {item!r}")

    print(f"saved: {OUT}")
    print(f"base paragraphs: {len(doc.paragraphs)}")
    print("existing content check: preserved")


if __name__ == "__main__":
    main()
