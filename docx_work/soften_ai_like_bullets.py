from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("output/doc/win_内容表述优化版.docx")
BODY_FONT = "Microsoft YaHei"
TEXT = RGBColor(0x22, 0x22, 0x22)


REPLACEMENTS = {
    "- 大学生与考研/考证人群：频繁截图学习资料（如政治提纲、代码片段、复习笔记），但难以系统化管理，形成“保存即遗忘”的恶性循环；": "从目标用户看，智存面向的是一类非常普遍的碎片化信息使用者。大学生和考研、考证人群会频繁保存政治提纲、代码片段、复习笔记等学习资料，却常常因为缺少整理入口而在真正需要时找不到；游戏玩家会积累攻略、地图和角色配装截图，但传统相册很难理解这些内容的语义；生活记录者会收藏旅游攻略、菜谱、购物清单和灵感速记，职场人士也会随手记录会议要点或项目想法。这些信息原本都有复用价值，却容易散落在相册、备忘录和聊天记录之间，逐渐变成难以回收的碎片。",
    "- 游戏玩家：保存大量攻略、地图、角色配装截图，查找困难，传统相册无法语义检索；": "",
    "- 生活记录者：收藏旅游攻略、菜谱、购物清单、灵感速记，散落在相册与备忘录中，无法沉淀为可复用资产；": "",
    "- 职场人士：快速记录会议要点、项目灵感，缺乏统一的知识回捞渠道，碎片信息难以转化为复利。": "",
    "- 用户每日会产生大量截图、速记，但是整理成本高、检索效率低，无差别淹没在日常相册与杂乱文件夹中；": "痛点的高频性也较为明确。用户每天都会产生截图、速记和临时收藏，但整理成本高、检索效率低，很多内容最终被无差别地淹没在相册或杂乱文件夹里。传统相册通常只能按时间、文件名或简单标签查找，笔记类应用又往往要求用户主动分类、主动编辑，这与碎片信息“看到就想保存、用完就想离开”的使用习惯并不匹配。更现实的问题是，学习和工作资料混入生活相册后，在展示私人照片时容易造成尴尬；如果依赖云端模型处理截图，聊天记录、证件信息和财务内容又会带来额外的隐私顾虑。",
    "- 传统相册仅支持按时间/文件名检索，笔记类APP需手工分类，违背“用完即走”的高效工具属性；": "",
    "- 现有解决方案存在“相册污染”问题：工作学习资产与生活影像混杂，打开相册向亲友展示时尴尬频发；": "",
    "- 云端AI处理方案存在隐私隐忧：截图常含聊天记录、证件信息、财务账单，用户不敢上传。": "",
    "- 日均截图次数：普通用户2-5次，重度用户可达10次以上；": "从使用频率看，普通用户每天可能产生数次截图，重度用户的截图数量还会更高；而搜索和回顾这些内容的需求，也会在每周复习、查攻略、找菜谱、整理灵感时反复出现。这说明智存面对的不是低频尝鲜场景，而是日常数字生活中持续发生的轻量刚需。",
    "- 搜索/回顾频率：每周至少3-5次（如找菜谱、查攻略、复习笔记、回看灵感）；": "",
    "- 交互无界化：全局悬浮球+MediaProjection技术，打破应用壁垒，实现“伴随式”无感捕获，不中断用户心流；": "对比现有方案，智存的优势并不只是多做了一个截图入口，而是把捕获、理解、沉淀和寻回放在同一条使用链路里。全局悬浮球与屏幕捕获能力让用户不必在多个应用之间来回切换，遇到高价值信息时可以直接保存，尽量减少对当前任务的打断。进入应用后，长截图、横屏图片和纯文本笔记会被统一整理成知识卡片，并通过智能文件夹和瀑布流进行呈现，使不同形态的内容能够放在同一个知识空间里被比较、浏览和再次使用。",
    "- 管理异构化：全异构瀑布流+智能文件夹，将长截图、横屏图片、纯文本统一抽象为“知识卡片”，图文跨越格式鸿沟和谐共存；": "",
    "- 检索语义化：双路召回（关键词+向量语义搜索）+首创“状态保活”机制，支持模糊意图检索（如搜“怎么去风龙废墟”找到“塞尔达地图”），并在结果列表与详情页之间零延迟无缝穿梭；": "在检索层面，智存同时保留关键词匹配和语义理解能力，既能处理明确的标签搜索，也能响应“怎么去风龙废墟”这类更接近日常表达的模糊意图。结果列表与详情页之间保留搜索状态和浏览位置，也让用户在反复筛选时不必重新开始。隐私方面，方案优先采用端侧数据沙盒和本地处理流程，把OCR、标签提取、向量索引等能力尽量放在设备内完成，截图不进入公共相册，也不默认上传云端，从而降低敏感信息外泄的风险，并让工作学习资产与生活影像保持清晰边界。",
    "- 安全端侧化：端侧数据沙盒+100%本地AI处理（OCR、标签提取、向量索引全部离线完成），截图不上云，工作资产与生活影像彻底分流，隐私零风险。": "",
}


def set_rfonts(run, font_name=BODY_FONT):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def apply_body_style(paragraph):
    paragraph.style = "Normal"
    paragraph.alignment = 3  # WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.32
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0.74)
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        set_rfonts(run, BODY_FONT)
        run.font.size = Pt(10.5)
        run.font.bold = False
        run.font.color.rgb = TEXT


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main():
    candidates = list(Path("output/doc").glob("win_*.docx"))
    source = max(candidates, key=lambda item: item.stat().st_mtime)
    doc = Document(str(source))
    original_paragraph_count = len(doc.paragraphs)

    changed = 0
    to_remove = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in REPLACEMENTS:
            replacement = REPLACEMENTS[text]
            if replacement:
                paragraph.text = replacement
                apply_body_style(paragraph)
            else:
                to_remove.append(paragraph)
            changed += 1

    for paragraph in to_remove:
        remove_paragraph(paragraph)

    remaining_dash = [p.text.strip() for p in doc.paragraphs if p.text.strip().startswith("- ")]
    if remaining_dash:
        raise RuntimeError("Dash-list paragraphs remain:\n" + "\n".join(remaining_dash))

    doc.save(str(OUT))

    final_doc = Document(str(OUT))
    print(f"source: {source}")
    print(f"saved: {OUT}")
    print(f"rewritten paragraphs: {changed}")
    print(f"removed list-only paragraphs: {len(to_remove)}")
    print(f"paragraph count: {original_paragraph_count} -> {len(final_doc.paragraphs)}")
    print("dash-list check: passed")


if __name__ == "__main__":
    main()
