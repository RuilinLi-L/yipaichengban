from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree


PATH = Path("output/doc/win_字体统一黑色版.docx")
FONT_NAME = "Microsoft YaHei"
BLACK = "000000"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W = NS["w"]


def qn(local):
    return f"{{{W}}}{local}"


def ensure_child(parent, local):
    child = parent.find(qn(local))
    if child is None:
        child = etree.Element(qn(local))
        parent.append(child)
    return child


def normalize_rpr(rpr):
    rfonts = ensure_child(rpr, "rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(attr), FONT_NAME)
    for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(attr), None)

    color = ensure_child(rpr, "color")
    color.set(qn("val"), BLACK)
    for attr in ("themeColor", "themeTint", "themeShade"):
        color.attrib.pop(qn(attr), None)


def patch_xml(data):
    root = etree.fromstring(data)
    changed = False
    for rpr in root.xpath(".//w:rPr", namespaces=NS):
        normalize_rpr(rpr)
        changed = True
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True), changed


def text_signature(path):
    from docx import Document

    doc = Document(str(path))
    items = [("p", p.text) for p in doc.paragraphs]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                items.append((f"t{ti}r{ri}c{ci}", cell.text))
    return items


def main():
    before = text_signature(PATH)
    patched = 0
    tmp_path = PATH.with_name(PATH.stem + ".tmp.docx")
    with ZipFile(PATH, "r") as zin:
        with ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                data = zin.read(info.filename)
                if info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    try:
                        data, changed = patch_xml(data)
                        patched += int(changed)
                    except etree.XMLSyntaxError:
                        pass
                zout.writestr(info, data)

    tmp_path.replace(PATH)
    after = text_signature(PATH)
    if before != after:
        raise RuntimeError("Text content changed during OOXML font/color cleanup.")

    print(f"saved: {PATH}")
    print(f"patched xml files: {patched}")
    print("content check: unchanged")


if __name__ == "__main__":
    main()
