from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC = Path("docx_work/win_original.docx")
OUT = Path("docx_work/win_layout_polished.docx")

BODY_FONT = "Microsoft YaHei"
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
SUB_ACCENT = RGBColor(0x2F, 0x68, 0x92)
TEXT = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
GRID = "A7B9C9"
HEADER_FILL = "D9EAF7"

H1_TEXTS = {
    "团队介绍",
    "作品设计理念",
    "作品原型设计",
    "创新点",
    "前景评估",
}

H2_TEXTS = {
    "①团队定位：",
    "②成员分工：",
    "③协作方式：",
    "④团队优势：",
    "①产品定位：",
    "②设计哲学：",
    "③用户体验目标：",
    "原型总体说明",
    "核心流程设计",
    "演示路径",
    "①交互无界化：全局悬浮与无感捕获的“伴随式”体验",
    "②管理异构化：从物理隔离走向“全异构瀑布流”资产重组",
    "③检索语义化：双路召回与首创“状态保活”寻回体验",
    "④安全端侧化：绝对隐私的本地“数据沙盒”",
    "1. 用户需求程度：",
    "2. 市场欢迎程度：具有独特优势、符合行业发展趋势",
}

H3_TEXTS = {
    "①目标用户群体广泛",
    "②痛点明确且高频发生",
    "③使用场景真实且高频",
    "①独特优势",
    "②符合行业发展趋势",
    "③市场欢迎度预判",
}


def set_rfonts(element, font_name):
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(key), font_name)


def configure_style(style, font, size_pt, bold=False, color=None):
    style.font.name = font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color
    set_rfonts(style.element, font)


def ensure_paragraph_style(doc, name):
    try:
        return doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        return style


def set_style_outline_level(style, level):
    ppr = style.element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(level))


def remove_num_pr(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is not None:
        ppr.remove(num_pr)


def set_keep_next(paragraph, keep=True):
    paragraph.paragraph_format.keep_with_next = keep
    paragraph.paragraph_format.keep_together = keep


def format_runs(paragraph, size_pt, bold=None, color=TEXT):
    for run in paragraph.runs:
        run.font.name = BODY_FONT
        set_rfonts(run._element, BODY_FONT)
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.bold = bold
        if color is not None:
            run.font.color.rgb = color


def add_or_update_cell_child(parent, child_name):
    child = parent.find(qn(child_name))
    if child is None:
        child = OxmlElement(child_name)
        parent.append(child)
    return child


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = add_or_update_cell_child(tc_pr, "w:tcMar")
    for side, value in {
        "w:top": top,
        "w:start": start,
        "w:bottom": bottom,
        "w:end": end,
    }.items():
        node = mar.find(qn(side))
        if node is None:
            node = OxmlElement(side)
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = add_or_update_cell_child(tc_pr, "w:shd")
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=GRID, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = add_or_update_cell_child(tc_pr, "w:tcBorders")
    for side in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
        edge = borders.find(qn(side))
        if edge is None:
            edge = OxmlElement(side)
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def text_signature(doc):
    sig = [("p", p.text) for p in doc.paragraphs]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                sig.append((f"t{ti}r{ri}c{ci}", cell.text))
    return sig


def classify_paragraph(paragraph):
    text = paragraph.text.strip()
    if not text:
        return "blank"
    if text in H1_TEXTS:
        return "h1"
    if text in H2_TEXTS:
        return "h2"
    if text in H3_TEXTS:
        return "h3"
    if text.startswith("- "):
        return "bullet"
    if "\n" in paragraph.text and len(text.split("\n", 1)[0]) <= 26:
        return "detail"
    if re.match(r"^[①②③④⑤⑥⑦⑧⑨⑩]", text):
        return "numbered_body"
    return "body"


def apply_paragraph_format(paragraph, kind):
    remove_num_pr(paragraph)
    pf = paragraph.paragraph_format
    pf.keep_together = False
    pf.keep_with_next = False
    pf.widow_control = True

    if kind == "blank":
        paragraph.style = "Normal"
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1
        format_runs(paragraph, 5, bold=False, color=TEXT)
        return

    if kind == "h1":
        paragraph.style = "Heading 1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(16)
        pf.space_after = Pt(8)
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.12
        set_keep_next(paragraph)
        format_runs(paragraph, 18, bold=True, color=ACCENT)
        return

    if kind == "h2":
        paragraph.style = "Heading 2"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(11)
        pf.space_after = Pt(5)
        pf.left_indent = Cm(0)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.15
        set_keep_next(paragraph)
        format_runs(paragraph, 13, bold=True, color=SUB_ACCENT)
        return

    if kind == "h3":
        paragraph.style = "Heading 3"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(7)
        pf.space_after = Pt(3)
        pf.left_indent = Cm(0.35)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.15
        set_keep_next(paragraph)
        format_runs(paragraph, 11.5, bold=True, color=SUB_ACCENT)
        return

    if kind == "bullet":
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.left_indent = Cm(0.82)
        pf.first_line_indent = Cm(-0.32)
        pf.line_spacing = 1.28
        format_runs(paragraph, 10.5, bold=False, color=TEXT)
        return

    if kind == "numbered_body":
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(2)
        pf.space_after = Pt(5)
        pf.left_indent = Cm(0.65)
        pf.first_line_indent = Cm(-0.35)
        pf.line_spacing = 1.3
        format_runs(paragraph, 10.5, bold=False, color=TEXT)
        return

    if kind == "detail":
        paragraph.style = "Normal"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.space_before = Pt(2)
        pf.space_after = Pt(6)
        pf.left_indent = Cm(0.55)
        pf.first_line_indent = Cm(0)
        pf.line_spacing = 1.28
        format_runs(paragraph, 10.3, bold=False, color=MUTED)
        return

    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(0)
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.32
    format_runs(paragraph, 10.5, bold=False, color=TEXT)


def format_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)
    set_repeat_header(table.rows[0])

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.18
                for run in p.runs:
                    run.font.name = BODY_FONT
                    set_rfonts(run._element, BODY_FONT)
                    run.font.size = Pt(9.3 if row_idx else 10)
                    run.bold = row_idx == 0
                    run.font.color.rgb = TEXT if row_idx else ACCENT


def main():
    doc = Document(str(SRC))
    before = text_signature(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.25)
    section.right_margin = Cm(2.25)

    normal_style = doc.styles["Normal"]
    heading1 = ensure_paragraph_style(doc, "Heading 1")
    heading2 = ensure_paragraph_style(doc, "Heading 2")
    heading3 = ensure_paragraph_style(doc, "Heading 3")

    configure_style(normal_style, BODY_FONT, 10.5, False, TEXT)
    configure_style(heading1, BODY_FONT, 18, True, ACCENT)
    configure_style(heading2, BODY_FONT, 13, True, SUB_ACCENT)
    configure_style(heading3, BODY_FONT, 11.5, True, SUB_ACCENT)
    set_style_outline_level(heading1, 0)
    set_style_outline_level(heading2, 1)
    set_style_outline_level(heading3, 2)

    normal_pf = normal_style.paragraph_format
    normal_pf.line_spacing = 1.32
    normal_pf.space_after = Pt(6)
    normal_pf.first_line_indent = Cm(0.74)

    for paragraph in doc.paragraphs:
        apply_paragraph_format(paragraph, classify_paragraph(paragraph))

    if len(doc.tables) >= 1:
        format_table(doc.tables[0], [3.2, 8.0, 4.9])
    if len(doc.tables) >= 2:
        format_table(doc.tables[1], [3.0, 4.1, 9.0])

    after = text_signature(doc)
    if before != after:
        raise RuntimeError("Text content changed during formatting; aborting save.")

    doc.save(str(OUT))

    reread = Document(str(OUT))
    if before != text_signature(reread):
        raise RuntimeError("Saved document text does not match source text.")

    print(f"saved: {OUT}")
    print(f"paragraphs: {len(doc.paragraphs)}, tables: {len(doc.tables)}")
    print("text check: unchanged")


if __name__ == "__main__":
    main()
